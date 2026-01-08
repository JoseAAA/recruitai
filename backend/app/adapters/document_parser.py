"""
Document Parser Adapter
Handles CV/Resume file parsing using PyMuPDF and python-docx.
"""
import hashlib
import logging
import re
import tempfile
from pathlib import Path
from typing import Tuple

logger = logging.getLogger(__name__)


class DocumentParsingError(Exception):
    """Raised when document parsing fails."""
    pass


class DocumentParser:
    """
    Parser for CV/Resume documents (PDF, DOCX).
    Uses PyMuPDF for PDFs and python-docx for Word documents.
    """
    
    def __init__(self):
        pass
    
    async def parse_file(
        self,
        file_path: str,
        detect_hidden_text: bool = True
    ) -> Tuple[str, dict]:
        """
        Parse a document file and extract text content.
        
        Args:
            file_path: Path to the file
            detect_hidden_text: Whether to check for hidden text manipulation
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        path = Path(file_path)
        
        if not path.exists():
            raise DocumentParsingError(f"File not found: {file_path}")
        
        suffix = path.suffix.lower()
        
        if suffix not in ['.pdf', '.docx', '.doc', '.txt']:
            raise DocumentParsingError(f"Unsupported file type: {suffix}")
        
        try:
            if suffix == '.pdf':
                text = self._parse_pdf(path)
            elif suffix in ['.docx', '.doc']:
                text = self._parse_docx(path)
            elif suffix == '.txt':
                text = path.read_text(encoding='utf-8', errors='ignore')
            else:
                raise DocumentParsingError(f"Unsupported: {suffix}")
            
            # Clean the text
            text = self.clean_text(text)
            
            # Calculate file hash
            file_hash = self._calculate_hash(path)
            
            metadata = {
                "file_path": str(path),
                "file_name": path.name,
                "file_size": path.stat().st_size,
                "file_hash": file_hash,
                "total_characters": len(text),
            }
            
            return text, metadata
            
        except Exception as e:
            logger.error(f"Failed to parse document {file_path}: {e}")
            raise DocumentParsingError(f"Parsing failed: {str(e)}")
    
    def _parse_pdf(self, path: Path) -> str:
        """Parse PDF using PyMuPDF (fitz)."""
        import fitz  # PyMuPDF
        
        text_parts = []
        
        with fitz.open(str(path)) as doc:
            for page in doc:
                text_parts.append(page.get_text())
        
        return "\n".join(text_parts)
    
    def _parse_docx(self, path: Path) -> str:
        """Parse DOCX using python-docx."""
        from docx import Document
        
        doc = Document(str(path))
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Also get text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        return "\n".join(text_parts)
    
    async def parse_bytes(
        self,
        content: bytes,
        filename: str
    ) -> Tuple[str, dict]:
        """
        Parse document from bytes content.
        
        Args:
            content: File content as bytes
            filename: Original filename (for extension detection)
            
        Returns:
            Tuple of (extracted_text, metadata)
        """
        suffix = Path(filename).suffix
        
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(content)
            tmp_path = tmp.name
        
        try:
            text, metadata = await self.parse_file(tmp_path, detect_hidden_text=True)
            metadata["original_filename"] = filename
            return text, metadata
        finally:
            # Cleanup temp file
            Path(tmp_path).unlink(missing_ok=True)
    
    @staticmethod
    def _calculate_hash(file_path: Path) -> str:
        """Calculate SHA-256 hash of file content."""
        sha256 = hashlib.sha256()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                sha256.update(chunk)
        return sha256.hexdigest()
    
    @staticmethod
    def clean_text(text: str) -> str:
        """Clean and normalize extracted text."""
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        
        # Remove common artifacts
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        
        return text.strip()
